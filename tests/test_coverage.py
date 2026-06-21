"""Format coverage: images (content_hash + agent record), DOCX, and XLSX."""
import pytest


@pytest.fixture(autouse=True)
def _isolated_home(tmp_path, monkeypatch):
    monkeypatch.setenv("DNR_HOME", str(tmp_path / "dnrhome"))


def _mkpng(path):
    from PIL import Image

    im = Image.new("RGB", (40, 30))
    for x in range(40):
        for y in range(30):
            im.putpixel((x, y), ((x * 6) % 256, (y * 8) % 256, 120))
    im.save(path)


def test_image_content_hash_stable(tmp_path):
    from dnr import hashing

    p = tmp_path / "a.png"
    _mkpng(p)
    assert hashing.content_hash(p) == hashing.content_hash(p)
    assert hashing.content_hash(p).startswith("sha256:")


def test_image_record_and_search(tmp_path):
    from dnr import embed, index, ingest

    folder = tmp_path / "img"
    folder.mkdir()
    p = folder / "chart.png"
    _mkpng(p)
    rec = ingest.record_supplied(p, "Bar chart: Q4 revenue 1,200,000 KRW by region", "vision", "claude-opus-4-vision")
    assert rec["provenance"]["method"] == "vision"
    assert embed.extract(p) is None  # default db-only keeps image bytes untouched
    assert index.db_only_record(folder, p) == rec
    index.scan(folder)
    assert index.query_match(folder, "revenue") == ["chart.png"]


def test_image_ingest_directs_to_record(tmp_path):
    from dnr import cli

    p = tmp_path / "x.png"
    _mkpng(p)
    assert cli.main(["ingest", str(p)]) == 1  # clean 'use dnr record', no crash


def test_docx_ingest_local(tmp_path):
    import docx

    from dnr import embed, hashing, index, ingest

    folder = tmp_path / "d"
    folder.mkdir()
    p = folder / "memo.docx"
    d = docx.Document()
    d.add_paragraph("Quarterly memo: contract renewal pending approval")
    d.save(str(p))
    rec = ingest.ingest(p)
    assert rec["provenance"]["transcriber"] == "python-docx"
    assert rec["content_hash"] == hashing.content_hash(p)
    assert embed.extract(p) is None  # docx has no in-file carrier yet -> db-only, file untouched
    assert index.db_only_record(folder, p) == rec
    assert index.query_match(folder, "renewal") == ["memo.docx"]  # db-only -> queryable


def test_docx_bytes_unchanged_by_ingest(tmp_path):
    import docx

    from dnr import hashing, ingest

    p = tmp_path / "m.docx"
    d = docx.Document()
    d.add_paragraph("body text")
    d.save(str(p))
    h0 = hashing.content_hash(p)
    ingest.ingest(p)  # writes a db-only record; the .docx zip is not touched
    assert hashing.content_hash(p) == h0


def test_xlsx_ingest_local(tmp_path):
    from openpyxl import Workbook

    from dnr import embed, hashing, index, ingest

    folder = tmp_path / "s"
    folder.mkdir()
    p = folder / "ledger.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Payments"
    ws.append(["Party", "Amount", "Status"])
    ws.append(["Acme", 1200000, "paid"])
    wb.save(str(p))

    rec = ingest.ingest(p)
    assert rec["provenance"]["transcriber"] == "openpyxl"
    assert "Payments" in rec["transcript"]["text"]
    assert "Acme\t1200000\tpaid" in rec["transcript"]["text"]
    assert rec["content_hash"] == hashing.content_hash(p)
    assert embed.extract(p) is None  # xlsx has no in-file carrier yet -> db-only
    assert index.db_only_record(folder, p) == rec
    assert index.query_match(folder, "Acme") == ["ledger.xlsx"]


def test_backfill_folder_ingests_local_and_lists_agent_needed(tmp_path):
    import docx
    from openpyxl import Workbook

    from dnr import index, ingest

    folder = tmp_path / "bf"
    folder.mkdir()

    pdf = folder / "contract.pdf"
    from fpdf import FPDF

    p = FPDF()
    p.add_page()
    p.set_font("Helvetica", size=12)
    p.cell(0, 8, "contract renewal damages")
    p.output(str(pdf))

    d = docx.Document()
    d.add_paragraph("docx approval memo")
    d.save(str(folder / "memo.docx"))

    wb = Workbook()
    ws = wb.active
    ws.append(["Party", "Amount"])
    ws.append(["Acme", 42])
    wb.save(str(folder / "ledger.xlsx"))

    _mkpng(folder / "scan.png")

    stats = ingest.backfill(folder)
    assert {x["path"] for x in stats["ingested"]} == {"contract.pdf", "memo.docx", "ledger.xlsx"}
    assert stats["agent_needed"] == [{"path": "scan.png", "reason": "needs agent/vision transcript"}]
    assert stats["errors"] == []
    assert index.query_match(folder, "Acme") == ["ledger.xlsx"]
    assert index.query_match(folder, "approval") == ["memo.docx"]


def test_db_only_record_removed_when_source_changes(tmp_path):
    import docx

    from dnr import index, ingest

    folder = tmp_path / "fresh"
    folder.mkdir()
    p = folder / "memo.docx"
    d = docx.Document()
    d.add_paragraph("old unique needle")
    d.save(str(p))

    ingest.ingest(p)
    index.scan(folder)
    assert index.query_match(folder, "needle") == ["memo.docx"]

    d = docx.Document()
    d.add_paragraph("new body without the old keyword")
    d.save(str(p))

    assert ingest.read_cached(p) is None
    stats = index.scan(folder)
    assert stats["removed"] == 1
    assert index.query_match(folder, "needle") == []
